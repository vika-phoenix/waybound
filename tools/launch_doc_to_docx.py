# -*- coding: utf-8 -*-
"""
Turn the launch-readiness page into a Word document.

Built from the same HTML rather than retyped, so the two cannot drift apart.
Only the tags that page actually uses are handled — this is a converter for one
document, not a general one, and pretending otherwise would be a lot of code
guarding against markup that is never going to appear.

    pip install python-docx
    python tools/launch_doc_to_docx.py         Kavkazland-Launch-Readiness.html Kavkazland-Launch-Readiness.docx

Edit the HTML, re-run this, commit both. The .docx in the repo is generated —
editing it directly means the next run silently throws your changes away.
"""
import io
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

SRC, OUT = sys.argv[1], sys.argv[2]

INK   = RGBColor(0x14, 0x20, 0x2A)
MUTE  = RGBColor(0x5E, 0x6F, 0x7A)
FAINT = RGBColor(0x8A, 0x99, 0xA2)
TEAL  = RGBColor(0x1C, 0x6F, 0x6A)
RUST  = RGBColor(0xA6, 0x3A, 0x2B)
AMBER = RGBColor(0x8A, 0x64, 0x20)
GREEN = RGBColor(0x2C, 0x7A, 0x56)

CHIP_COLOUR = {'stop': RUST, 'wait': AMBER, 'ok': GREEN, 'go': TEAL, '': MUTE}

SANS, MONO = 'Segoe UI', 'Consolas'


# ── parse ────────────────────────────────────────────────────────────────────

class Reader(HTMLParser):
    """Collapse the page into an ordered list of blocks with inline runs."""

    BLOCKS = {'p', 'h1', 'h2', 'h3', 'li', 'td', 'th'}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self.stack = []          # open tags, with their class attribute
        self.runs = None         # (text, bold, italic, mono) while inside a block
        self.table = None
        self.row = None
        self.skip = 0

    # -- helpers
    def _classes(self, tag=None):
        for t, cls in reversed(self.stack):
            if tag is None or t == tag:
                return cls
        return set()

    def _in(self, tag):
        return any(t == tag for t, _ in self.stack)

    def _ancestor_class(self, name):
        return any(name in cls for _, cls in self.stack)

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        cls = set((a.get('class') or '').split())
        if tag in ('style', 'title', 'script'):
            self.skip += 1
            return
        self.stack.append((tag, cls))

        if tag == 'table':
            self.table = {'kind': 'table', 'head': [], 'rows': []}
        elif tag == 'tr':
            self.row = []
        elif tag in self.BLOCKS:
            self.runs = []
        elif tag == 'div' and self._ancestor_class('tally'):
            # The count strip is spans inside bare divs, so it is invisible to
            # the block rule. It is the summary of the whole page — worth the
            # one special case rather than losing it.
            self.runs = []
        elif tag == 'br' and self.runs is not None:
            self.runs.append(('\n', False, False, False))

    def handle_endtag(self, tag):
        if tag in ('style', 'title', 'script'):
            self.skip = max(0, self.skip - 1)
            return
        cls = set()
        for i in range(len(self.stack) - 1, -1, -1):
            if self.stack[i][0] == tag:
                cls = self.stack[i][1]
                del self.stack[i:]
                break

        if tag in ('td', 'th'):
            if self.row is not None:
                self.row.append(self._take(cls))
            self.runs = None
        elif tag == 'tr':
            if self.row:
                inside_head = self._in('thead')
                (self.table['head'] if inside_head else self.table['rows']).append(self.row)
            self.row = None
        elif tag == 'table':
            self.blocks.append(self.table)
            self.table = None
        elif tag == 'div' and self.runs is not None and self._ancestor_class('tally'):
            block = self._take(cls)
            if block['runs']:
                block['kind'] = 'tally'
                self.blocks.append(block)
            self.runs = None
        elif tag in self.BLOCKS:
            runs = self._take(cls)
            if runs['runs']:
                runs['kind'] = 'li' if tag == 'li' else tag
                runs['ordered'] = self._ancestor_class('steps')
                self.blocks.append(runs)
            self.runs = None

    def _take(self, cls):
        runs = self.runs or []
        # Trim the whitespace HTML indentation leaves at the edges.
        while runs and not runs[0][0].strip():
            runs.pop(0)
        while runs and not runs[-1][0].strip():
            runs.pop()
        return {'kind': 'p', 'runs': runs, 'cls': cls}

    def handle_data(self, data):
        if self.skip or self.runs is None:
            return
        text = re.sub(r'\s+', ' ', data)
        if not text:
            return
        self.runs.append((
            text,
            self._in('strong') or self._in('b'),
            self._in('em') or self._in('i'),
            self._in('code') or self._in('kbd'),
        ))


# ── emit ─────────────────────────────────────────────────────────────────────

def shade(el, hex_fill):
    tc = OxmlElement('w:shd')
    tc.set(qn('w:val'), 'clear')
    tc.set(qn('w:fill'), hex_fill)
    el.get_or_add_tcPr().append(tc) if el.tag.endswith('tc') else el.append(tc)


def write_runs(par, runs, colour=None, size=None, bold_all=False):
    for text, bold, italic, mono in runs:
        r = par.add_run(text)
        r.bold = bold or bold_all
        r.italic = italic
        r.font.name = MONO if mono else SANS
        r.font.size = Pt(size or 10.5) if not mono else Pt((size or 10.5) - 0.5)
        r.font.color.rgb = colour or INK


def plain(runs):
    return ''.join(t for t, *_ in runs).strip()


def build(blocks, doc):
    pending_ref = None
    pending_chips = []

    tally = []

    def flush_tally():
        if not tally:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        for i, entry in enumerate(tally):
            if i:
                sep = p.add_run('   ·   ')
                sep.font.name, sep.font.size = SANS, Pt(9)
                sep.font.color.rgb = FAINT
            r = p.add_run(entry)
            r.font.name, r.font.size, r.bold = MONO, Pt(9.5), True
            r.font.color.rgb = INK
        tally.clear()

    for b in blocks:
        cls, kind = b.get('cls', set()), b['kind']

        if kind == 'tally':
            # "1" and "blocks everything" arrive as two spans in one div.
            parts = [t.strip() for t, *_ in b['runs'] if t.strip()]
            tally.append(' '.join(parts))
            continue
        if tally and kind != 'tally':
            flush_tally()

        if kind == 'table':
            emit_table(doc, b)
            continue

        text = plain(b['runs'])
        if not text:
            continue

        # The ref (P1) and the chips (BLOCKS LAUNCH) are separate spans beside
        # the heading in HTML. In Word they read better folded into its line.
        if 'ref' in cls:
            pending_ref = text
            continue
        if 'chip' in cls:
            pending_chips.append((text, next((c for c in ('stop', 'wait', 'ok', 'go') if c in cls), '')))
            continue

        if kind == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            write_runs(p, b['runs'], size=25, bold_all=True)
            pending_chips = []
            continue

        if kind == 'h2':
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph()
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(10), Pt(3)
            write_runs(p, b['runs'], size=16, bold_all=True)
            for label, tone in pending_chips:
                c = p.add_run('   ' + label.upper())
                c.font.name, c.font.size, c.bold = MONO, Pt(8), True
                c.font.color.rgb = CHIP_COLOUR[tone]
            pending_chips = []
            continue

        if kind == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(14), Pt(4)
            if pending_ref:
                r = p.add_run(pending_ref + '  ')
                r.font.name, r.font.size, r.bold = MONO, Pt(10), True
                r.font.color.rgb = FAINT
                pending_ref = None
            write_runs(p, b['runs'], size=12.5, bold_all=True)
            for label, tone in pending_chips:
                c = p.add_run('   ' + label.upper())
                c.font.name, c.font.size, c.bold = MONO, Pt(8), True
                c.font.color.rgb = CHIP_COLOUR[tone]
            pending_chips = []
            continue

        if kind == 'li':
            p = doc.add_paragraph(style='List Number' if b.get('ordered') else 'List Bullet')
            p.paragraph_format.space_after = Pt(3)
            write_runs(p, b['runs'])
            continue

        # ── paragraph variants
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)

        if 'eyebrow' in cls:
            r = p.add_run(text.upper())
            r.font.name, r.font.size, r.bold = MONO, Pt(8.5), True
            r.font.color.rgb = TEAL
        elif 'standfirst' in cls:
            write_runs(p, b['runs'], colour=MUTE, size=12)
        elif 'asof' in cls:
            r = p.add_run(text)
            r.font.name, r.font.size = MONO, Pt(8.5)
            r.font.color.rgb = FAINT
        elif 'sec-note' in cls:
            write_runs(p, b['runs'], colour=MUTE, size=10)
        elif 'verdict' in cls:
            p.paragraph_format.space_before = Pt(11)
            r = p.add_run(text.upper())
            r.font.name, r.font.size, r.bold = MONO, Pt(8.5), True
            r.font.color.rgb = TEAL
        elif 'l' in cls or 'n' in cls:      # the tally strip
            write_runs(p, b['runs'], colour=MUTE, size=10)
        else:
            in_note = 'note' in cls or any('note' in c for c in [cls])
            write_runs(p, b['runs'])
            if in_note:
                p.paragraph_format.left_indent = Inches(0.22)


def emit_table(doc, spec):
    head, rows = spec['head'], spec['rows']
    width = max([len(r) for r in head + rows] or [1])
    t = doc.add_table(rows=0, cols=width)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for hrow in head:
        cells = t.add_row().cells
        for i, cell in enumerate(hrow[:width]):
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            r = para.add_run(plain(cell['runs']).upper())
            r.font.name, r.font.size, r.bold = MONO, Pt(7.5), True
            r.font.color.rgb = FAINT
            shade(cells[i]._tc, 'EDF0EF')

    for row in rows:
        cells = t.add_row().cells
        for i, cell in enumerate(row[:width]):
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            write_runs(para, cell['runs'], size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    html = io.open(SRC, encoding='utf-8').read()
    reader = Reader()
    reader.feed(html)

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name, normal.font.size = SANS, Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    build(reader.blocks, doc)
    doc.save(OUT)
    print('wrote %s — %d blocks' % (OUT, len(reader.blocks)))


main()
